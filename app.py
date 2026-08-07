from __future__ import annotations

import csv
import json
import os
import queue
import sys
import threading
import traceback
from datetime import datetime
from pathlib import Path

if sys.platform == "win32":
    os.environ.setdefault("QT_OPENGL", "software")
    os.environ.setdefault("QT_QUICK_BACKEND", "software")
    os.environ.setdefault("QSG_RHI_BACKEND", "software")

from PySide6.QtCore import QEvent, QPoint, QRect, QSize, Qt, QTimer, QUrl, Signal
from PySide6.QtGui import (
    QAction,
    QColor,
    QCloseEvent,
    QFont,
    QFontDatabase,
    QFontMetrics,
    QIcon,
    QMouseEvent,
    QPainter,
    QPainterPath,
    QPen,
    QPixmap,
    QTextCursor,
)
from PySide6.QtWidgets import (
    QApplication,
    QButtonGroup,
    QCheckBox,
    QColorDialog,
    QComboBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFormLayout,
    QFrame,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QScrollArea,
    QSlider,
    QSizeGrip,
    QSpinBox,
    QSplitter,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from xfyun_client import XfyunInterpreter, set_manuscript_cache_db

APP_NAME = "ísmolar 同声传译 · v1.9.7 简洁界面版"
APP_VERSION = "1.9.7"

TOKENS = {
    "bg": "#F6F9FD",
    "bg_alt": "#F3F7FC",
    "deep_blue": "#2F5EA8",
    "blue": "#1768D5",
    "green": "#08AD91",
    "light_blue": "#EAF2FC",
    "text": "#243044",
    "muted": "#66758A",
    "border": "#D9E3F0",
    "white": "#FFFFFF",
    "danger": "#C43D4B",
    "disabled": "#E9EDF3",
}

DEFAULT_CONFIG = {
    "app_id": "",
    "api_key": "",
    "api_secret": "",
    "deepseek_api_key": "",
    "deepseek_model": "deepseek-v4-flash",
    "translation_direction": "auto",
    "play_tts": False,
    "show_chinese": True,
    "overlay_display_mode": "translation",
    "translation_interval_mode": "adaptive",
    "background_opacity": 45,
    "subtitle_font": "Microsoft YaHei",
    "chinese_size": 34,
    "english_size": 32,
    "subtitle_color": "#FFFFFF",
    "outline_color": "#000000",
    "outline_enabled": True,
    "subtitle_split_ratio": 50,
    "overlay_width": 980,
    "overlay_height": 330,
    "source_language": "zh",
    "glossary_entries": [],
    "glossary_file_name": "",
    "reference_text": "",
    "reference_file_name": "",
}


def resource_path(relative: str) -> Path:
    base = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    return base / relative


def config_dir() -> Path:
    if sys.platform == "darwin":
        return Path.home() / "Library" / "Application Support" / "ismolar-interpreter"
    if sys.platform == "win32":
        return Path(os.environ.get("APPDATA", Path.home())) / "ismolar-interpreter"
    return Path.home() / ".config" / "ismolar-interpreter"


CONFIG_DIR = config_dir()
CONFIG_FILE = CONFIG_DIR / "config.json"
LOG_DIR = CONFIG_DIR / "logs"
LOG_FILE = LOG_DIR / "app.log"
RECORD_DIR = Path.home() / "Documents" / "ísmolar同声传译记录"


def append_log(text: str) -> None:
    try:
        LOG_DIR.mkdir(parents=True, exist_ok=True)
        with LOG_FILE.open("a", encoding="utf-8") as fp:
            fp.write(f"[{datetime.now():%Y-%m-%d %H:%M:%S}] {text}\n")
    except Exception:
        pass


def load_config() -> dict:
    merged = dict(DEFAULT_CONFIG)
    if CONFIG_FILE.exists():
        try:
            data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                merged.update(data)
        except Exception as exc:
            append_log(f"读取配置失败：{exc}")
    return merged


def save_config(config: dict) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    CONFIG_FILE.write_text(
        json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    try:
        os.chmod(CONFIG_FILE, 0o600)
    except OSError:
        pass


MAX_GLOSSARY_ENTRIES = 500
MAX_REFERENCE_CHARS = 40000


def _read_text_file(path: Path) -> str:
    last_error: Exception | None = None
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5"):
        try:
            return path.read_text(encoding=encoding)
        except Exception as exc:
            last_error = exc
    raise RuntimeError(f"无法读取文本文件：{last_error}")


def _clean_glossary_rows(rows) -> list[list[str]]:
    entries: list[list[str]] = []
    seen: set[tuple[str, str]] = set()
    for row in rows:
        values = [str(value).strip() if value is not None else "" for value in row]
        if len(values) < 2:
            continue
        source, target = values[0], values[1]
        if not source or not target:
            continue
        if source.lower() in {"原词", "原文", "source", "term", "英文", "english"} and target.lower() in {
            "译文", "固定译法", "target", "translation", "中文", "chinese"
        }:
            continue
        key = (source.casefold(), target.casefold())
        if key in seen:
            continue
        seen.add(key)
        entries.append([source, target])
        if len(entries) >= MAX_GLOSSARY_ENTRIES:
            break
    return entries


def load_glossary_file(path: Path) -> list[list[str]]:
    suffix = path.suffix.lower()
    rows: list[list[str]] = []

    if suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("缺少 openpyxl，重新运行一键制作脚本即可安装。") from exc
        workbook = load_workbook(path, read_only=True, data_only=True)
        sheet = workbook.active
        rows = [[cell for cell in row[:2]] for row in sheet.iter_rows(values_only=True)]
    elif suffix in {".csv", ".tsv"}:
        text = _read_text_file(path)
        delimiter = "\t" if suffix == ".tsv" else ","
        rows = [list(row) for row in csv.reader(text.splitlines(), delimiter=delimiter)]
    elif suffix == ".txt":
        text = _read_text_file(path)
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#"):
                continue
            pair = None
            for delimiter in ("\t", "=>", "→", "->", "=", "：", ":"):
                if delimiter in line:
                    left, right = line.split(delimiter, 1)
                    pair = [left, right]
                    break
            if pair is None and "," in line:
                left, right = line.split(",", 1)
                pair = [left, right]
            if pair:
                rows.append(pair)
    else:
        raise RuntimeError("固定翻译支持 TXT、CSV、TSV、XLSX。")

    entries = _clean_glossary_rows(rows)
    if not entries:
        raise RuntimeError(
            "没有识别到有效的固定译法。每行需要两列，例如：BIOEFFECT,蓓欧菲"
        )
    return entries


def load_reference_file(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".txt":
        text = _read_text_file(path)
    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("缺少 python-docx，重新运行一键制作脚本即可安装。") from exc
        document = Document(path)
        text = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
    elif suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("缺少 openpyxl，重新运行一键制作脚本即可安装。") from exc
        workbook = load_workbook(path, read_only=True, data_only=True)
        lines: list[str] = []
        for row in workbook.active.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                lines.append("\t".join(values))
        text = "\n".join(lines)
    elif suffix in {".csv", ".tsv"}:
        raw = _read_text_file(path)
        delimiter = "\t" if suffix == ".tsv" else ","
        text = "\n".join("\t".join(row) for row in csv.reader(raw.splitlines(), delimiter=delimiter))
    else:
        raise RuntimeError("参考稿件支持 TXT、DOCX、CSV、TSV、XLSX。")

    text = text.strip()
    if not text:
        raise RuntimeError("导入的参考稿件没有文字内容。")
    return text[:MAX_REFERENCE_CHARS]


class WaveIndicator(QWidget):
    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setFixedSize(42, 22)
        self._running = False
        self._phase = 0
        self._timer = QTimer(self)
        self._timer.timeout.connect(self._tick)

    def set_running(self, running: bool) -> None:
        self._running = running
        if running:
            self._timer.start(180)
        else:
            self._timer.stop()
        self.update()

    def _tick(self) -> None:
        self._phase = (self._phase + 1) % 4
        self.update()

    def paintEvent(self, _event) -> None:  # type: ignore[override]
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        heights = [6, 12, 18, 10, 15]
        for i, base in enumerate(heights):
            h = base if not self._running else max(5, base + ((i + self._phase) % 3 - 1) * 4)
            x = 4 + i * 7
            wave_pen = QPen(
                QColor(TOKENS["blue"] if self._running else TOKENS["border"])
            )
            wave_pen.setWidthF(3.0)
            wave_pen.setCapStyle(Qt.PenCapStyle.RoundCap)
            painter.setPen(wave_pen)
            painter.drawLine(x, 11 - h // 2, x, 11 + h // 2)


class CredentialDialog(QDialog):
    def __init__(self, config: dict, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setWindowTitle("接口设置")
        self.setMinimumWidth(580)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(16)

        title = QLabel("接口设置")
        title.setObjectName("dialogTitle")
        layout.addWidget(title)
        note = QLabel("密钥只保存在当前电脑，不会写入 GitHub 仓库。")
        note.setObjectName("mutedLabel")
        layout.addWidget(note)

        form = QFormLayout()
        form.setHorizontalSpacing(18)
        form.setVerticalSpacing(14)
        self.app_id = QLineEdit(str(config.get("app_id", "")))
        self.api_key = QLineEdit(str(config.get("api_key", "")))
        self.api_secret = QLineEdit(str(config.get("api_secret", "")))
        self.deepseek_api_key = QLineEdit(str(config.get("deepseek_api_key", "")))
        self.deepseek_model = QLineEdit(str(config.get("deepseek_model", "deepseek-v4-flash")))
        for widget in (self.api_key, self.api_secret, self.deepseek_api_key):
            widget.setEchoMode(QLineEdit.EchoMode.Password)
        form.addRow("讯飞 APPID", self.app_id)
        form.addRow("讯飞 APIKey", self.api_key)
        form.addRow("讯飞 APISecret", self.api_secret)
        form.addRow("DeepSeek API Key", self.deepseek_api_key)
        form.addRow("DeepSeek 模型", self.deepseek_model)
        layout.addLayout(form)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Save
            | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def values(self) -> dict:
        return {
            "app_id": self.app_id.text().strip(),
            "api_key": self.api_key.text().strip(),
            "api_secret": self.api_secret.text().strip(),
            "deepseek_api_key": self.deepseek_api_key.text().strip(),
            "deepseek_model": self.deepseek_model.text().strip() or "deepseek-v4-flash",
        }


class TranslationTimingDialog(QDialog):
    OPTIONS = (
        ("自适应 2–5 秒", "adaptive"),
        ("固定 2 秒", "2"),
        ("固定 3 秒", "3"),
        ("固定 4 秒", "4"),
        ("固定 5 秒", "5"),
        ("等待句末", "sentence"),
    )

    def __init__(self, current: str, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setWindowTitle("翻译时间")
        self.setMinimumWidth(360)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(22, 22, 22, 22)
        layout.setSpacing(14)
        title = QLabel("翻译刷新时间")
        title.setObjectName("dialogTitle")
        layout.addWidget(title)
        note = QLabel(
            "自适应模式会根据当前句子长度，在 2–5 秒内更新临时译文；"
            "新结果会覆盖上一版临时译文，主界面仍保留完整记录。"
        )
        note.setWordWrap(True)
        note.setObjectName("mutedLabel")
        layout.addWidget(note)
        self.combo = QComboBox()
        for label, value in self.OPTIONS:
            self.combo.addItem(label, value)
        index = self.combo.findData(current)
        self.combo.setCurrentIndex(index if index >= 0 else 0)
        layout.addWidget(self.combo)
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Save
            | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def value(self) -> str:
        return str(self.combo.currentData() or "adaptive")


class OverlaySettingsDialog(QDialog):
    def __init__(self, settings: dict, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setWindowTitle("浮窗设置")
        self.setMinimumWidth(520)
        self._text_color = QColor(str(settings.get("subtitle_color", "#FFFFFF")))
        self._outline_color = QColor(str(settings.get("outline_color", "#000000")))

        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(16)
        title = QLabel("浮窗设置")
        title.setObjectName("dialogTitle")
        layout.addWidget(title)

        form = QFormLayout()
        form.setHorizontalSpacing(18)
        form.setVerticalSpacing(13)

        self.direction = QComboBox()
        self.direction.addItem("自动识别中英文", "auto")
        self.direction.addItem("中译英", "zh_en")
        self.direction.addItem("英译中", "en_zh")
        direction_index = self.direction.findData(settings.get("translation_direction", "zh_en"))
        self.direction.setCurrentIndex(max(0, direction_index))
        form.addRow("翻译方向", self.direction)

        self.tts = QCheckBox("启用翻译发音")
        self.tts.setChecked(bool(settings.get("play_tts", False)))
        form.addRow("语音", self.tts)

        self.opacity = QSlider(Qt.Orientation.Horizontal)
        self.opacity.setRange(0, 90)
        self.opacity.setValue(int(settings.get("background_opacity", 45)))
        form.addRow("背景透明度", self.opacity)

        self.font = QComboBox()
        families = QFontDatabase.families()
        preferred_fonts = (
            "Microsoft YaHei", "微软雅黑", "PingFang SC", "SF Pro Display", "Arial"
        )
        for family in preferred_fonts:
            if family in families:
                self.font.addItem(family)
        if self.font.count() == 0:
            self.font.addItems(families[:100])
        font_index = self.font.findText(str(settings.get("subtitle_font", "Microsoft YaHei")))
        if font_index >= 0:
            self.font.setCurrentIndex(font_index)
        form.addRow("字幕字体", self.font)

        self.cn_size = QSpinBox()
        self.cn_size.setRange(12, 96)
        self.cn_size.setValue(int(settings.get("chinese_size", 34)))
        self.cn_size.setSuffix(" pt")
        form.addRow("中文字号", self.cn_size)

        self.en_size = QSpinBox()
        self.en_size.setRange(12, 96)
        self.en_size.setValue(int(settings.get("english_size", 32)))
        self.en_size.setSuffix(" pt")
        form.addRow("英文字号", self.en_size)

        self.text_color_button = QPushButton()
        self.text_color_button.clicked.connect(self._choose_text_color)
        self._update_color_button(self.text_color_button, self._text_color, "字色")
        form.addRow("字幕颜色", self.text_color_button)

        self.outline_color_button = QPushButton()
        self.outline_color_button.clicked.connect(self._choose_outline_color)
        self._update_color_button(self.outline_color_button, self._outline_color, "描边色")
        form.addRow("描边颜色", self.outline_color_button)

        self.outline = QCheckBox("启用字幕描边")
        self.outline.setChecked(bool(settings.get("outline_enabled", True)))
        form.addRow("描边", self.outline)
        layout.addLayout(form)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Save
            | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    @staticmethod
    def _update_color_button(button: QPushButton, color: QColor, label: str) -> None:
        button.setText(f"{label}  {color.name().upper()}")
        button.setStyleSheet(
            f"QPushButton {{ border-left: 22px solid {color.name()}; }}"
        )

    def _choose_text_color(self) -> None:
        color = QColorDialog.getColor(self._text_color, self, "选择字幕颜色")
        if color.isValid():
            self._text_color = color
            self._update_color_button(self.text_color_button, color, "字色")

    def _choose_outline_color(self) -> None:
        color = QColorDialog.getColor(self._outline_color, self, "选择描边颜色")
        if color.isValid():
            self._outline_color = color
            self._update_color_button(self.outline_color_button, color, "描边色")

    def values(self) -> dict:
        return {
            "translation_direction": str(self.direction.currentData() or "zh_en"),
            "play_tts": self.tts.isChecked(),
            "background_opacity": self.opacity.value(),
            "subtitle_font": self.font.currentText() or "Microsoft YaHei",
            "chinese_size": self.cn_size.value(),
            "english_size": self.en_size.value(),
            "subtitle_color": self._text_color.name().upper(),
            "outline_color": self._outline_color.name().upper(),
            "outline_enabled": self.outline.isChecked(),
        }


class OutlinedTextWidget(QWidget):
    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self._text = ""
        self._family = "Microsoft YaHei"
        self._size = 34
        self._color = QColor("#FFFFFF")
        self._outline_color = QColor("#000000")
        self._outline_enabled = True

        # 字幕安全区：最后一行始终与底边、状态栏和缩放把手保持距离。
        self._horizontal_padding = 14
        self._top_padding = 8
        self._bottom_safe_padding = 20

        self.setObjectName("subtitleText")
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)
        self.setAutoFillBackground(False)
        self.setMinimumHeight(48)

    def set_text(self, text: str) -> None:
        self._text = text
        self.update()

    def set_style_options(
        self,
        family: str,
        size: int,
        color: QColor,
        outline_color: QColor,
        outline_enabled: bool,
    ) -> None:
        self._family = family
        self._size = int(size)
        self._color = QColor(color)
        self._outline_color = QColor(outline_color)
        self._outline_enabled = bool(outline_enabled)

        # 字号越大，底部安全距离同步增加；最大值限制避免浪费过多空间。
        self._bottom_safe_padding = max(
            20,
            min(54, round(self._size * 0.42)),
        )
        self.update()

    def _text_draw_rect(self, font: QFont) -> tuple[QRect, QRect]:
        content_rect = self.rect().adjusted(
            self._horizontal_padding,
            self._top_padding,
            -self._horizontal_padding,
            -self._bottom_safe_padding,
        )
        if content_rect.width() <= 2 or content_rect.height() <= 2:
            return content_rect, content_rect

        flags = Qt.AlignmentFlag.AlignLeft | Qt.TextFlag.TextWordWrap
        metrics = QFontMetrics(font)

        # 用足够高的测量区域计算完整文字高度，然后把整段文字锚定到底部。
        # 文字过多时只会从顶部裁切，最后一行不会被底边遮挡。
        measure_rect = QRect(0, 0, content_rect.width(), 1_000_000)
        required = metrics.boundingRect(measure_rect, flags, self._text)
        required_height = max(metrics.lineSpacing(), required.height())

        draw_rect = QRect(
            content_rect.left(),
            content_rect.bottom() - required_height + 1,
            content_rect.width(),
            required_height,
        )
        return content_rect, draw_rect

    def paintEvent(self, _event) -> None:  # type: ignore[override]
        if not self._text:
            return

        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)

        font = QFont(self._family, self._size)
        font.setHintingPreference(QFont.HintingPreference.PreferFullHinting)
        painter.setFont(font)

        content_rect, draw_rect = self._text_draw_rect(font)
        if content_rect.width() <= 2 or content_rect.height() <= 2:
            return

        flags = Qt.AlignmentFlag.AlignLeft | Qt.TextFlag.TextWordWrap
        painter.save()
        painter.setClipRect(content_rect)

        if self._outline_enabled:
            outline_pen = QPen(self._outline_color)
            outline_pen.setWidthF(4.0)
            outline_pen.setStyle(Qt.PenStyle.SolidLine)
            outline_pen.setCapStyle(Qt.PenCapStyle.RoundCap)
            outline_pen.setJoinStyle(Qt.PenJoinStyle.RoundJoin)
            painter.setPen(outline_pen)
            painter.drawText(draw_rect.translated(-1, 0), flags, self._text)
            painter.drawText(draw_rect.translated(1, 0), flags, self._text)
            painter.drawText(draw_rect.translated(0, -1), flags, self._text)
            painter.drawText(draw_rect.translated(0, 1), flags, self._text)

        painter.setPen(self._color)
        painter.drawText(draw_rect, flags, self._text)
        painter.restore()


class SubtitleOverlay(QMainWindow):
    stop_requested = Signal()
    direction_changed = Signal(str)
    tts_changed = Signal(bool)
    timing_changed = Signal(str)
    settings_changed = Signal(dict)

    DISPLAY_MODES = {
        "source": "原文",
        "translation": "译文",
        "bilingual": "双语",
    }

    def __init__(self, config: dict) -> None:
        super().__init__()
        self.setObjectName("subtitleOverlay")
        self.setWindowTitle(APP_NAME)
        self.setWindowFlags(
            Qt.WindowType.FramelessWindowHint
            | Qt.WindowType.WindowStaysOnTopHint
            | Qt.WindowType.Tool
        )
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)
        self.setAutoFillBackground(False)
        self.resize(
            max(300, int(config.get("overlay_width", 980))),
            max(56, int(config.get("overlay_height", 330))),
        )
        self.setMinimumSize(300, 56)
        self._drag_origin: QPoint | None = None
        self._drag_window_origin: QPoint | None = None
        self._press_global: QPoint | None = None
        self._full_chinese = ""
        self._full_english = ""
        self._source_transcript = ""
        self._translation_transcript = ""
        self._source_language = str(config.get("source_language", "zh"))
        self._direction = str(config.get("translation_direction", "zh_en"))
        self._display_mode = str(config.get("overlay_display_mode", "translation"))
        self._timing_mode = str(config.get("translation_interval_mode", "adaptive"))
        self._play_tts = bool(config.get("play_tts", False))
        self._opacity = int(config.get("background_opacity", 45))
        self._font_family = str(config.get("subtitle_font", "Microsoft YaHei"))
        self._cn_size_value = int(config.get("chinese_size", 34))
        self._en_size_value = int(config.get("english_size", 32))
        self._text_color = QColor(config.get("subtitle_color", "#FFFFFF"))
        self._outline_color = QColor(config.get("outline_color", "#000000"))
        self._outline_enabled = bool(config.get("outline_enabled", True))
        self._last_split_ratio = int(config.get("subtitle_split_ratio", 50))

        self.toolbar_hide_timer = QTimer(self)
        self.toolbar_hide_timer.setSingleShot(True)
        self.toolbar_hide_timer.timeout.connect(self._hide_toolbar)
        self._build_ui()
        self.apply_config(config)

    def _build_ui(self) -> None:
        central = QWidget()
        central.setObjectName("overlayRoot")
        central.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)
        central.setAutoFillBackground(False)
        self.setCentralWidget(central)
        outer = QHBoxLayout(central)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(6)

        self.side_toolbar = QFrame()
        self.side_toolbar.setObjectName("overlaySideToolbar")
        self.side_toolbar.setFixedWidth(64)
        side = QVBoxLayout(self.side_toolbar)
        side.setContentsMargins(6, 7, 6, 7)
        side.setSpacing(6)

        self.mode_group = QButtonGroup(self)
        self.mode_group.setExclusive(True)
        self.mode_buttons: dict[str, QPushButton] = {}
        for mode, label in self.DISPLAY_MODES.items():
            button = QPushButton(label)
            button.setObjectName("overlayModeButton")
            button.setCheckable(True)
            button.setFixedHeight(34)
            self.mode_group.addButton(button)
            button.clicked.connect(
                lambda _checked=False, selected=mode: self._set_display_mode(selected)
            )
            self.mode_buttons[mode] = button
            side.addWidget(button)

        side.addStretch(1)
        self.time_button = QPushButton("时间")
        self.time_button.setObjectName("overlayToolButton")
        self.time_button.setToolTip("设置 2–5 秒翻译刷新时间")
        self.time_button.clicked.connect(self._open_timing_dialog)
        side.addWidget(self.time_button)

        self.settings_button = QPushButton("设置")
        self.settings_button.setObjectName("overlayToolButton")
        self.settings_button.clicked.connect(self._open_settings_dialog)
        side.addWidget(self.settings_button)

        stop_button = QPushButton("停止")
        stop_button.setObjectName("overlayStop")
        stop_button.clicked.connect(self.stop_requested.emit)
        side.addWidget(stop_button)
        outer.addWidget(self.side_toolbar)

        self.container = QFrame()
        self.container.setObjectName("overlayContainer")
        layout = QVBoxLayout(self.container)
        layout.setContentsMargins(10, 8, 10, 10)
        layout.setSpacing(4)

        self.subtitle_splitter = QSplitter(Qt.Orientation.Vertical)
        self.subtitle_splitter.setObjectName("subtitleSplitter")
        self.subtitle_splitter.setChildrenCollapsible(False)
        self.subtitle_splitter.setOpaqueResize(True)
        self.subtitle_splitter.setHandleWidth(10)
        self.primary_subtitle = OutlinedTextWidget()
        self.secondary_subtitle = OutlinedTextWidget()
        self.primary_subtitle.setMinimumHeight(20)
        self.secondary_subtitle.setMinimumHeight(20)
        self.subtitle_splitter.addWidget(self.primary_subtitle)
        self.subtitle_splitter.addWidget(self.secondary_subtitle)
        self.subtitle_splitter.setStretchFactor(0, 1)
        self.subtitle_splitter.setStretchFactor(1, 1)
        self.subtitle_splitter.splitterMoved.connect(self._splitter_moved)
        layout.addWidget(self.subtitle_splitter, 1)

        bottom = QHBoxLayout()
        bottom.setContentsMargins(3, 4, 2, 2)
        bottom.setSpacing(6)
        self.compact_status = QLabel("")
        self.compact_status.setObjectName("overlayStatus")
        self.compact_status.setMinimumHeight(20)
        bottom.addWidget(self.compact_status)
        bottom.addStretch(1)
        self.resize_grip = QSizeGrip(self.container)
        self.resize_grip.setObjectName("overlayResizeGrip")
        self.resize_grip.setFixedSize(18, 18)
        bottom.addWidget(self.resize_grip)
        layout.addLayout(bottom)

        outer.addWidget(self.container, 1)
        for watched in (
            self.container,
            self.subtitle_splitter,
            self.primary_subtitle,
            self.secondary_subtitle,
        ):
            watched.installEventFilter(self)
        self.side_toolbar.hide()

    def apply_config(self, config: dict) -> None:
        self._direction = str(config.get("translation_direction", self._direction))
        self._display_mode = str(config.get("overlay_display_mode", self._display_mode))
        if self._display_mode not in self.DISPLAY_MODES:
            self._display_mode = "translation"
        self._timing_mode = str(config.get("translation_interval_mode", self._timing_mode))
        self._play_tts = bool(config.get("play_tts", self._play_tts))
        self._opacity = int(config.get("background_opacity", self._opacity))
        self._font_family = str(config.get("subtitle_font", self._font_family))
        self._cn_size_value = int(config.get("chinese_size", self._cn_size_value))
        self._en_size_value = int(config.get("english_size", self._en_size_value))
        self._text_color = QColor(str(config.get("subtitle_color", self._text_color.name())))
        self._outline_color = QColor(str(config.get("outline_color", self._outline_color.name())))
        self._outline_enabled = bool(config.get("outline_enabled", self._outline_enabled))
        self._last_split_ratio = max(
            5, min(95, int(config.get("subtitle_split_ratio", self._last_split_ratio)))
        )
        self._source_language = str(config.get("source_language", self._source_language))
        self._sync_mode_buttons()
        self._apply_style()
        self._refresh_display()

    def current_settings(self) -> dict:
        return {
            "translation_direction": self._direction,
            "overlay_display_mode": self._display_mode,
            "translation_interval_mode": self._timing_mode,
            "play_tts": self._play_tts,
            "show_chinese": self._display_mode != "translation" or self._target_language() == "zh",
            "background_opacity": self._opacity,
            "subtitle_font": self._font_family,
            "chinese_size": self._cn_size_value,
            "english_size": self._en_size_value,
            "subtitle_color": self._text_color.name().upper(),
            "outline_color": self._outline_color.name().upper(),
            "outline_enabled": self._outline_enabled,
            "subtitle_split_ratio": self._last_split_ratio,
            "overlay_width": self.width(),
            "overlay_height": self.height(),
            "source_language": self._source_language,
        }

    def set_direction(self, direction: str) -> None:
        self._direction = direction if direction in {"auto", "zh_en", "en_zh"} else "zh_en"
        self._refresh_display()

    def set_play_tts(self, enabled: bool) -> None:
        self._play_tts = bool(enabled)

    def set_timing_mode(self, mode: str) -> None:
        self._timing_mode = mode if mode in {"adaptive", "2", "3", "4", "5", "sentence"} else "adaptive"

    def set_subtitles(
        self,
        chinese: str,
        english: str,
        source_language: str | None = None,
        source_transcript: str = "",
        translation_transcript: str = "",
    ) -> None:
        self._full_chinese = chinese
        self._full_english = english
        self._source_transcript = source_transcript
        self._translation_transcript = translation_transcript
        if source_language in {"zh", "en"}:
            self._source_language = source_language
        self._refresh_display()

    def set_status(self, text: str) -> None:
        self.compact_status.setText(text)
        # 当尚未产生字幕时，同步刷新占位提示，避免透明浮窗看起来像空屏。
        if not self._full_chinese.strip() and not self._full_english.strip():
            self._refresh_display()

    def show_for_translation(self) -> None:
        # 启动时先显示工具栏和状态提示，数秒后再自动隐藏。
        # 这样即使尚未产生译文，也不会出现“整个屏幕什么都没有”的感觉。
        if not self.compact_status.text().strip():
            self.compact_status.setText("正在连接，请开始讲话…")
        self._refresh_display()
        self.side_toolbar.show()
        self.side_toolbar.raise_()
        self._schedule_toolbar_hide()
        self.show()
        self.raise_()
        self.activateWindow()

    def _source_language_for_display(self) -> str:
        if self._direction == "zh_en":
            return "zh"
        if self._direction == "en_zh":
            return "en"
        return "en" if self._source_language == "en" else "zh"

    def _target_language(self) -> str:
        return "zh" if self._source_language_for_display() == "en" else "en"

    def _text_for_language(self, language: str) -> str:
        return self._full_chinese if language == "zh" else self._full_english

    def _apply_widget_style(self, widget: OutlinedTextWidget, language: str) -> None:
        if language == "zh":
            size = self._cn_size_value
        elif language == "en":
            size = self._en_size_value
        else:
            # 自动混合模式使用固定字号，切换语言时不改变整个浮窗的视觉布局。
            size = max(self._cn_size_value, self._en_size_value)
        widget.set_style_options(
            self._font_family,
            size,
            self._text_color,
            self._outline_color,
            self._outline_enabled,
        )

    def _refresh_display(self) -> None:
        if self._direction == "auto":
            # 自动模式按句子保存“原文流”和“译文流”，不再依据最新识别语言
            # 重新解释整段历史内容。中文→英文、英文→中文的结果可在同一
            # 译文浮窗中按发生顺序连续显示。
            source_text = self._source_transcript.strip()
            target_text = self._translation_transcript.strip()
            source_style_language = "mixed"
            target_style_language = "mixed"
        else:
            source_language = self._source_language_for_display()
            target_language = "zh" if source_language == "en" else "en"
            source_text = self._text_for_language(source_language).strip()
            target_text = self._text_for_language(target_language).strip()
            source_style_language = source_language
            target_style_language = target_language

        waiting_source = "请开始讲话…"
        waiting_translation = "正在翻译…" if source_text else "请开始讲话…"

        if self._display_mode == "bilingual":
            self.primary_subtitle.set_text(source_text or waiting_source)
            self.secondary_subtitle.set_text(target_text or waiting_translation)
            self._apply_widget_style(self.primary_subtitle, source_style_language)
            self._apply_widget_style(self.secondary_subtitle, target_style_language)
            self.primary_subtitle.show()
            self.secondary_subtitle.show()
            self.subtitle_splitter.setHandleWidth(10)
            QTimer.singleShot(0, self._restore_split)
        else:
            if self._display_mode == "source":
                visible_text = source_text or waiting_source
                style_language = source_style_language
            else:
                visible_text = target_text or waiting_translation
                style_language = target_style_language
            self.primary_subtitle.set_text(visible_text)
            self._apply_widget_style(self.primary_subtitle, style_language)
            self.primary_subtitle.show()
            self.secondary_subtitle.hide()
            self.subtitle_splitter.setHandleWidth(0)
        self._sync_mode_buttons()

    def _sync_mode_buttons(self) -> None:
        for mode, button in self.mode_buttons.items():
            button.blockSignals(True)
            button.setChecked(mode == self._display_mode)
            button.blockSignals(False)

    def _set_display_mode(self, mode: str) -> None:
        if mode not in self.DISPLAY_MODES:
            return
        self._display_mode = mode
        self._refresh_display()
        self.settings_changed.emit(self.current_settings())
        self._schedule_toolbar_hide()

    def _open_timing_dialog(self) -> None:
        self.toolbar_hide_timer.stop()
        dialog = TranslationTimingDialog(self._timing_mode, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self._timing_mode = dialog.value()
            self.timing_changed.emit(self._timing_mode)
            self.settings_changed.emit(self.current_settings())
        self._schedule_toolbar_hide()

    def _open_settings_dialog(self) -> None:
        self.toolbar_hide_timer.stop()
        before_direction = self._direction
        before_tts = self._play_tts
        dialog = OverlaySettingsDialog(self.current_settings(), self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            values = dialog.values()
            self._direction = str(values["translation_direction"])
            self._play_tts = bool(values["play_tts"])
            self._opacity = int(values["background_opacity"])
            self._font_family = str(values["subtitle_font"])
            self._cn_size_value = int(values["chinese_size"])
            self._en_size_value = int(values["english_size"])
            self._text_color = QColor(str(values["subtitle_color"]))
            self._outline_color = QColor(str(values["outline_color"]))
            self._outline_enabled = bool(values["outline_enabled"])
            self._apply_style()
            self._refresh_display()
            if self._direction != before_direction:
                self.direction_changed.emit(self._direction)
            if self._play_tts != before_tts:
                self.tts_changed.emit(self._play_tts)
            self.settings_changed.emit(self.current_settings())
        self._schedule_toolbar_hide()

    def _apply_style(self) -> None:
        background_alpha = max(0, min(230, round(self._opacity * 2.55)))
        toolbar_alpha = 0 if self._opacity == 0 else max(90, min(235, background_alpha + 60))
        self.container.setStyleSheet(
            "QFrame#overlayContainer {"
            f"background-color: rgba(20, 32, 52, {background_alpha});"
            "border: 1px solid rgba(217, 227, 240, 65);"
            "border-radius: 12px;"
            "}"
            "QWidget#subtitleText, QSplitter#subtitleSplitter { background: transparent; }"
        )
        self.side_toolbar.setStyleSheet(
            "QFrame#overlaySideToolbar {"
            f"background-color: rgba(246, 249, 253, {toolbar_alpha});"
            "border: 1px solid rgba(217, 227, 240, 160);"
            "border-radius: 11px;"
            "}"
        )
        self._refresh_display()

    def _splitter_moved(self, _pos: int, _index: int) -> None:
        sizes = self.subtitle_splitter.sizes()
        if len(sizes) >= 2 and sum(sizes[:2]) > 0 and self._display_mode == "bilingual":
            self._last_split_ratio = max(
                5, min(95, round(sizes[0] * 100 / sum(sizes[:2])))
            )
            self.settings_changed.emit(self.current_settings())

    def _restore_split(self) -> None:
        if self._display_mode == "bilingual":
            self.subtitle_splitter.setSizes(
                [self._last_split_ratio, 100 - self._last_split_ratio]
            )

    def _reveal_toolbar(self) -> None:
        self.side_toolbar.show()
        self.side_toolbar.raise_()
        self._schedule_toolbar_hide()

    def _schedule_toolbar_hide(self) -> None:
        self.toolbar_hide_timer.start(4500)

    def _hide_toolbar(self) -> None:
        self.side_toolbar.hide()

    def eventFilter(self, watched, event) -> bool:  # type: ignore[override]
        if watched in (
            self.container,
            self.subtitle_splitter,
            self.primary_subtitle,
            self.secondary_subtitle,
        ):
            if event.type() == QEvent.Type.MouseButtonPress and isinstance(event, QMouseEvent):
                if event.button() == Qt.MouseButton.LeftButton:
                    self._drag_origin = event.globalPosition().toPoint()
                    self._drag_window_origin = self.pos()
                    self._press_global = event.globalPosition().toPoint()
            elif event.type() == QEvent.Type.MouseMove and isinstance(event, QMouseEvent):
                if self._drag_origin is not None and event.buttons() & Qt.MouseButton.LeftButton:
                    delta = event.globalPosition().toPoint() - self._drag_origin
                    if abs(delta.x()) + abs(delta.y()) > 5:
                        self.move(self._drag_window_origin + delta)
                        return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                release_pos = (
                    event.globalPosition().toPoint()
                    if isinstance(event, QMouseEvent)
                    else self._press_global
                )
                if self._press_global is not None and release_pos is not None:
                    delta = release_pos - self._press_global
                    if abs(delta.x()) + abs(delta.y()) <= 6:
                        self._reveal_toolbar()
                self._drag_origin = None
                self._drag_window_origin = None
                self._press_global = None
                self.settings_changed.emit(self.current_settings())
        return super().eventFilter(watched, event)


class BilingualDocumentWindow(QMainWindow):
    back_requested = Signal()

    def __init__(self) -> None:
        super().__init__()
        self.setObjectName("documentWindow")
        self.setWindowTitle("传译文档 · 中英对照")
        self.resize(1160, 820)
        self.setMinimumSize(760, 560)
        self._segments: list[dict[str, str]] = []
        self._chinese = ""
        self._english = ""
        self._started_at: datetime | None = None
        self._finished_at: datetime | None = None

        icon_path = resource_path("assets/app_icon.png")
        if icon_path.exists():
            self.setWindowIcon(QIcon(str(icon_path)))

        central = QWidget()
        self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setContentsMargins(30, 24, 30, 24)
        root.setSpacing(16)

        top = QHBoxLayout()
        title_box = QVBoxLayout()
        title = QLabel("传译文档")
        title.setObjectName("documentTitle")
        title_box.addWidget(title)
        self.meta_label = QLabel("中英对照记录")
        self.meta_label.setObjectName("documentMeta")
        title_box.addWidget(self.meta_label)
        top.addLayout(title_box)
        top.addStretch(1)

        save_button = QPushButton("保存文档")
        save_button.setObjectName("outlineButton")
        save_button.clicked.connect(self._save_document)
        top.addWidget(save_button)

        back_button = QPushButton("返回设置")
        back_button.setObjectName("primaryButton")
        back_button.clicked.connect(self._return_to_settings)
        top.addWidget(back_button)
        root.addLayout(top)

        column_header = QFrame()
        column_header.setObjectName("documentColumnHeader")
        column_layout = QHBoxLayout(column_header)
        column_layout.setContentsMargins(58, 10, 16, 10)
        cn_header = QLabel("中文")
        cn_header.setObjectName("documentColumnTitle")
        en_header = QLabel("English")
        en_header.setObjectName("documentColumnTitle")
        column_layout.addWidget(cn_header, 1)
        column_layout.addWidget(en_header, 1)
        root.addWidget(column_header)

        self.scroll = QScrollArea()
        self.scroll.setObjectName("documentScroll")
        self.scroll.setWidgetResizable(True)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.document_body = QWidget()
        self.document_body.setObjectName("documentBody")
        self.document_layout = QVBoxLayout(self.document_body)
        self.document_layout.setContentsMargins(4, 4, 4, 20)
        self.document_layout.setSpacing(12)
        self.scroll.setWidget(self.document_body)
        root.addWidget(self.scroll, 1)

    def set_document(
        self,
        segments: list[dict[str, str]],
        chinese: str,
        english: str,
        started_at: datetime | None,
        finished_at: datetime | None,
    ) -> None:
        self._segments = [dict(segment) for segment in segments]
        self._chinese = chinese.strip()
        self._english = english.strip()
        self._started_at = started_at
        self._finished_at = finished_at

        started = started_at.strftime("%Y-%m-%d %H:%M") if started_at else "—"
        duration = ""
        if started_at and finished_at:
            seconds = max(0, int((finished_at - started_at).total_seconds()))
            duration = f" · {seconds // 60} 分 {seconds % 60} 秒"
        completed = sum(
            1 for segment in self._segments
            if str(segment.get("translation", "")).strip()
        )
        self.meta_label.setText(
            f"{started}{duration} · {len(self._segments)} 段中英对照"
            f" · 已完成 {completed}/{len(self._segments)}"
        )
        self._rebuild_document()
        QTimer.singleShot(0, lambda: self.scroll.verticalScrollBar().setValue(0))

    @staticmethod
    def _clear_layout(layout) -> None:
        while layout.count():
            item = layout.takeAt(0)
            widget = item.widget()
            child_layout = item.layout()
            if widget is not None:
                widget.deleteLater()
            elif child_layout is not None:
                BilingualDocumentWindow._clear_layout(child_layout)

    def _rebuild_document(self) -> None:
        self._clear_layout(self.document_layout)
        if not self._segments:
            empty = QLabel("本次传译还没有形成可整理的中英对照内容。")
            empty.setObjectName("documentEmpty")
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.document_layout.addWidget(empty)
            self.document_layout.addStretch(1)
            return

        for index, segment in enumerate(self._segments, start=1):
            source = str(segment.get("source", "")).strip()
            translation = str(segment.get("translation", "")).strip()
            language = str(segment.get("source_language", "zh"))
            if language == "en":
                chinese, english = translation, source
                direction = "英 → 中"
            else:
                chinese, english = source, translation
                direction = "中 → 英"

            block = QFrame()
            block.setObjectName("documentBlock")
            block_layout = QHBoxLayout(block)
            block_layout.setContentsMargins(14, 14, 14, 14)
            block_layout.setSpacing(14)

            number = QLabel(f"{index:02d}")
            number.setObjectName("documentNumber")
            number.setFixedWidth(36)
            number.setAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignHCenter)
            number.setToolTip(direction)
            block_layout.addWidget(number)

            for language_title, text in (("中文", chinese), ("English", english)):
                column = QFrame()
                column.setObjectName("documentColumn")
                column_box = QVBoxLayout(column)
                column_box.setContentsMargins(14, 12, 14, 14)
                column_box.setSpacing(8)
                label = QLabel(language_title)
                label.setObjectName("documentLanguageLabel")
                column_box.addWidget(label)
                missing_translation = not text and (
                    (language == "en" and language_title == "中文")
                    or (language != "en" and language_title == "English")
                )
                body = QLabel("翻译未完成" if missing_translation else (text or "—"))
                if missing_translation:
                    body.setObjectName("documentMissing")
                if not missing_translation:
                    body.setObjectName("documentParagraph")
                body.setWordWrap(True)
                body.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
                body.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignTop)
                column_box.addWidget(body)
                block_layout.addWidget(column, 1)

            self.document_layout.addWidget(block)
        self.document_layout.addStretch(1)

    def _return_to_settings(self) -> None:
        self.hide()
        self.back_requested.emit()

    def closeEvent(self, event: QCloseEvent) -> None:  # type: ignore[override]
        self.hide()
        self.back_requested.emit()
        event.ignore()

    def _save_document(self) -> None:
        if not self._segments and not self._chinese and not self._english:
            QMessageBox.information(self, "没有内容", "当前没有可保存的传译内容。")
            return
        RECORD_DIR.mkdir(parents=True, exist_ok=True)
        default = RECORD_DIR / f"中英对照传译文档_{datetime.now():%Y-%m-%d_%H-%M-%S}.docx"
        selected, selected_filter = QFileDialog.getSaveFileName(
            self,
            "保存中英对照文档",
            str(default),
            "Word 文档 (*.docx);;文本文件 (*.txt)",
        )
        if not selected:
            return
        path = Path(selected)
        try:
            if path.suffix.lower() == ".txt" or "文本" in selected_filter:
                if path.suffix.lower() != ".txt":
                    path = path.with_suffix(".txt")
                lines = ["ísmolar 同声传译 · 中英对照文档", ""]
                for index, segment in enumerate(self._segments, start=1):
                    source = str(segment.get("source", "")).strip()
                    translation = str(segment.get("translation", "")).strip()
                    if str(segment.get("source_language", "zh")) == "en":
                        chinese, english = translation, source
                    else:
                        chinese, english = source, translation
                    lines.extend([
                        f"{index:02d}",
                        f"中文：{chinese}",
                        f"English: {english}",
                        "",
                    ])
                path.write_text("\n".join(lines), encoding="utf-8")
            else:
                if path.suffix.lower() != ".docx":
                    path = path.with_suffix(".docx")
                from docx import Document
                from docx.enum.table import WD_TABLE_ALIGNMENT

                document = Document()
                document.add_heading("ísmolar 同声传译 · 中英对照文档", level=0)
                document.add_paragraph(self.meta_label.text())
                table = document.add_table(rows=1, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                table.rows[0].cells[0].text = "中文"
                table.rows[0].cells[1].text = "English"
                for segment in self._segments:
                    source = str(segment.get("source", "")).strip()
                    translation = str(segment.get("translation", "")).strip()
                    if str(segment.get("source_language", "zh")) == "en":
                        chinese, english = translation, source
                    else:
                        chinese, english = source, translation
                    cells = table.add_row().cells
                    cells[0].text = chinese
                    cells[1].text = english
                document.save(path)
        except Exception as exc:
            QMessageBox.critical(self, "保存失败", str(exc))
            return
        QMessageBox.information(self, "保存成功", f"文档已保存到：\n{path}")


class MainWindow(QMainWindow):
    def __init__(self) -> None:
        super().__init__()
        self.config = load_config()
        self.events: queue.Queue[dict] = queue.Queue()
        self.client: XfyunInterpreter | None = None
        self.client_generation = 0
        self.overlay: SubtitleOverlay | None = None
        self.result_window: BilingualDocumentWindow | None = None
        self.devices: list[tuple[int, str]] = []
        self.session_segments: list[dict[str, str]] = []
        self.session_chinese = ""
        self.session_english = ""
        self.session_started_at: datetime | None = None
        self.session_finished_at: datetime | None = None
        self._ending_session = False
        self._service_retry_count = 0
        self._finalize_thread: threading.Thread | None = None

        # 配置应用到 UI 时，复选框和字号控件可能触发保存回调。
        # 因此必须先创建 save_timer，再构建并恢复界面状态。
        self.save_timer = QTimer(self)
        self.save_timer.setSingleShot(True)
        self.save_timer.timeout.connect(lambda: save_config(self.config))

        self._build_ui()
        self._apply_config_to_ui()

        self.poll_timer = QTimer(self)
        self.poll_timer.timeout.connect(self._poll_events)
        self.poll_timer.start(80)
        QTimer.singleShot(300, self._refresh_devices)

        if not self._xfyun_ready():
            QTimer.singleShot(500, self._open_settings)

    def _build_ui(self) -> None:
        self.setWindowTitle(APP_NAME)
        self.resize(780, 660)
        self.setMinimumSize(660, 600)
        icon_path = resource_path("assets/app_icon.png")
        if icon_path.exists():
            self.setWindowIcon(QIcon(str(icon_path)))

        central = QWidget()
        self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setContentsMargins(42, 26, 42, 26)
        root.setSpacing(0)

        # 顶部：品牌 + 版本号，不再重复大标题和说明段落。
        header = QHBoxLayout()
        header.setSpacing(0)
        brand_logo = QLabel()
        brand_logo.setObjectName("brandLogo")
        logo_path = resource_path("assets/brand_logo.png")
        logo_pixmap = QPixmap(str(logo_path))
        if logo_pixmap.isNull():
            brand_logo.setText("ísmolar")
        else:
            logo_pixmap.setDevicePixelRatio(2.0)
            brand_logo.setPixmap(logo_pixmap)
            brand_logo.setFixedSize(130, 34)
        header.addWidget(brand_logo)
        header.addStretch(1)
        version_label = QLabel(f"v{APP_VERSION}")
        version_label.setObjectName("versionLabel")
        header.addWidget(version_label)
        root.addLayout(header)
        root.addSpacing(26)

        # 状态：一条大状态文字，替代原来的说明段落。
        self.status_label = QLabel("准备就绪")
        self.status_label.setObjectName("heroStatus")
        root.addWidget(self.status_label)
        root.addSpacing(8)

        # 翻译方向：三个分段按钮。
        direction_row = QHBoxLayout()
        direction_row.setSpacing(6)
        self.direction_group = QButtonGroup(self)
        self.direction_group.setExclusive(True)
        self.auto_button = QPushButton("自动")
        self.zh_en_button = QPushButton("中译英")
        self.en_zh_button = QPushButton("英译中")
        for button, direction in (
            (self.auto_button, "auto"),
            (self.zh_en_button, "zh_en"),
            (self.en_zh_button, "en_zh"),
        ):
            button.setCheckable(True)
            button.setObjectName("directionButton")
            self.direction_group.addButton(button)
            button.clicked.connect(lambda _checked=False, d=direction: self._change_direction(d))
            direction_row.addWidget(button)
        direction_row.addStretch(1)
        self.direction_note = QLabel()
        self.direction_note.setObjectName("flowNote")
        direction_row.addWidget(self.direction_note)
        root.addLayout(direction_row)
        root.addSpacing(20)

        # 开始 / 停止：主操作按钮。
        session_row = QHBoxLayout()
        session_row.setSpacing(10)
        self.start_button = QPushButton("开始传译")
        self.start_button.setObjectName("startButton")
        self.start_button.clicked.connect(self._start)
        session_row.addWidget(self.start_button, 1)
        self.stop_button = QPushButton("停止")
        self.stop_button.setObjectName("stopButton")
        self.stop_button.setEnabled(False)
        self.stop_button.clicked.connect(self._stop)
        session_row.addWidget(self.stop_button)
        self.wave = WaveIndicator()
        session_row.addWidget(self.wave)
        root.addLayout(session_row)
        root.addSpacing(28)

        # 设置区：轻量行 + 细分隔线，替代通栏卡片。
        settings_title = QLabel("设置")
        settings_title.setObjectName("settingsTitle")
        root.addWidget(settings_title)
        root.addSpacing(8)

        # 接口
        row = QHBoxLayout()
        row.setSpacing(10)
        row.addWidget(self._settings_label("接口"))
        self.api_dot = QLabel("●")
        self.api_dot.setObjectName("statusDot")
        row.addWidget(self.api_dot)
        self.api_state = QLabel("接口尚未配置")
        self.api_state.setObjectName("apiState")
        row.addWidget(self.api_state)
        row.addStretch(1)
        settings_button = QPushButton("接口设置")
        settings_button.setObjectName("outlineButton")
        settings_button.clicked.connect(self._open_settings)
        row.addWidget(settings_button)
        root.addLayout(row)
        root.addWidget(self._hairline())
        root.addSpacing(8)

        # 麦克风
        row = QHBoxLayout()
        row.setSpacing(10)
        row.addWidget(self._settings_label("麦克风"))
        self.device_combo = QComboBox()
        self.device_combo.setObjectName("settingCombo")
        self.device_combo.addItem("系统默认麦克风（推荐）", None)
        row.addWidget(self.device_combo, 1)
        refresh_button = QPushButton("刷新设备")
        refresh_button.setObjectName("outlineButton")
        refresh_button.clicked.connect(self._refresh_devices)
        row.addWidget(refresh_button)
        root.addLayout(row)
        root.addWidget(self._hairline())
        root.addSpacing(8)

        # 翻译发音
        row = QHBoxLayout()
        row.setSpacing(10)
        row.addWidget(self._settings_label("翻译发音"))
        row.addStretch(1)
        self.play_tts = QCheckBox("中英双向")
        self.play_tts.toggled.connect(self._tts_changed)
        row.addWidget(self.play_tts)
        root.addLayout(row)
        root.addWidget(self._hairline())
        root.addSpacing(8)

        # 浮窗显示 / 翻译刷新
        row = QHBoxLayout()
        row.setSpacing(10)
        row.addWidget(self._settings_label("浮窗"))
        self.overlay_mode_combo = QComboBox()
        self.overlay_mode_combo.setObjectName("settingCombo")
        self.overlay_mode_combo.addItem("仅译文", "translation")
        self.overlay_mode_combo.addItem("仅原文", "source")
        self.overlay_mode_combo.addItem("双语", "bilingual")
        self.overlay_mode_combo.setFixedWidth(118)
        self.overlay_mode_combo.currentIndexChanged.connect(self._overlay_mode_changed)
        row.addWidget(self.overlay_mode_combo)
        row.addSpacing(18)
        refresh_label = QLabel("翻译刷新")
        refresh_label.setObjectName("settingsLabel")
        row.addWidget(refresh_label)
        self.timing_combo = QComboBox()
        self.timing_combo.setObjectName("settingCombo")
        for label, value in TranslationTimingDialog.OPTIONS:
            self.timing_combo.addItem(label, value)
        self.timing_combo.setFixedWidth(150)
        self.timing_combo.currentIndexChanged.connect(self._translation_timing_changed)
        row.addWidget(self.timing_combo)
        row.addStretch(1)
        root.addLayout(row)
        root.addWidget(self._hairline())
        root.addSpacing(8)

        # 翻译资料
        row = QHBoxLayout()
        row.setSpacing(10)
        row.addWidget(self._settings_label("翻译资料"))
        self.materials_status = QLabel("未导入固定翻译或参考稿件")
        self.materials_status.setObjectName("mutedLabel")
        row.addWidget(self.materials_status, 1)
        root.addLayout(row)
        row = QHBoxLayout()
        row.setSpacing(8)
        row.addSpacing(66)
        glossary_button = QPushButton("上传固定译法")
        glossary_button.setObjectName("outlineButton")
        glossary_button.setToolTip("导入 TXT、CSV、TSV 或 XLSX 术语表")
        glossary_button.clicked.connect(self._import_glossary)
        row.addWidget(glossary_button)
        reference_button = QPushButton("导入参考稿件")
        reference_button.setObjectName("outlineButton")
        reference_button.setToolTip("导入 TXT、DOCX、CSV、TSV 或 XLSX 稿件")
        reference_button.clicked.connect(self._import_reference)
        row.addWidget(reference_button)
        clear_materials_button = QPushButton("清除资料")
        clear_materials_button.setObjectName("ghostButton")
        clear_materials_button.clicked.connect(self._clear_translation_materials)
        row.addWidget(clear_materials_button)
        row.addStretch(1)
        root.addLayout(row)

        root.addStretch(1)

    @staticmethod
    def _settings_label(text: str) -> QLabel:
        label = QLabel(text)
        label.setObjectName("settingsLabel")
        return label

    @staticmethod
    def _hairline() -> QFrame:
        line = QFrame()
        line.setObjectName("hairline")
        line.setFixedHeight(1)
        return line

    def _ensure_overlay(self) -> SubtitleOverlay:
        if self.overlay is None:
            self.overlay = SubtitleOverlay(self.config)
            self.overlay.stop_requested.connect(self._stop)
            self.overlay.direction_changed.connect(self._change_direction)
            self.overlay.tts_changed.connect(self._tts_changed)
            self.overlay.timing_changed.connect(self._translation_timing_changed)
            self.overlay.settings_changed.connect(self._overlay_settings_changed)
        return self.overlay

    def _ensure_result_window(self) -> BilingualDocumentWindow:
        if self.result_window is None:
            self.result_window = BilingualDocumentWindow()
            self.result_window.back_requested.connect(self._return_to_settings)
        return self.result_window

    def _return_to_settings(self) -> None:
        if self.result_window is not None:
            self.result_window.hide()
        self.show()
        self.raise_()
        self.activateWindow()
        self.status_label.setText("准备就绪")

    def _show_result_document(self) -> None:
        self.session_finished_at = self.session_finished_at or datetime.now()
        result = self._ensure_result_window()
        result.set_document(
            self.session_segments,
            self.session_chinese,
            self.session_english,
            self.session_started_at,
            self.session_finished_at,
        )
        result.show()
        result.raise_()
        result.activateWindow()

    def _abort_session_to_settings(self, title: str, message: str) -> None:
        self.client = None
        self._ending_session = False
        if self.overlay is not None:
            self.overlay.hide()
        self.start_button.setEnabled(True)
        self.stop_button.setEnabled(False)
        self.wave.set_running(False)
        self.show()
        self.raise_()
        self.activateWindow()
        self.status_label.setText(title)
        QMessageBox.critical(self, title, message)

    def _xfyun_ready(self) -> bool:
        return all(str(self.config.get(key, "")).strip() for key in ("app_id", "api_key", "api_secret"))

    def _apply_config_to_ui(self) -> None:
        direction = str(self.config.get("translation_direction", "zh_en"))
        if direction not in {"auto", "zh_en", "en_zh"}:
            direction = "zh_en"
        self.auto_button.setChecked(direction == "auto")
        self.zh_en_button.setChecked(direction == "zh_en")
        self.en_zh_button.setChecked(direction == "en_zh")
        self.play_tts.setChecked(bool(self.config.get("play_tts", False)))
        mode_index = self.overlay_mode_combo.findData(
            self.config.get("overlay_display_mode", "translation")
        )
        self.overlay_mode_combo.setCurrentIndex(max(0, mode_index))
        timing_index = self.timing_combo.findData(
            self.config.get("translation_interval_mode", "adaptive")
        )
        self.timing_combo.setCurrentIndex(max(0, timing_index))
        self._sync_direction_note(direction)
        self._update_materials_status()
        if self.overlay is not None:
            self.overlay.apply_config(self.config)
        self._refresh_api_state()

    def _refresh_api_state(self) -> None:
        xfyun = self._xfyun_ready()
        deepseek = bool(str(self.config.get("deepseek_api_key", "")).strip())
        if xfyun and deepseek:
            self.api_state.setText("讯飞同传 & DeepSeek 接口已配置")
            self.api_state.setProperty("state", "ok")
            self.api_dot.setProperty("state", "ok")
        elif xfyun:
            self.api_state.setText("讯飞同传已配置 · DeepSeek 未配置")
            self.api_state.setProperty("state", "partial")
            self.api_dot.setProperty("state", "partial")
        else:
            self.api_state.setText("接口尚未配置")
            self.api_state.setProperty("state", "error")
            self.api_dot.setProperty("state", "error")
        for widget in (self.api_state, self.api_dot):
            widget.style().unpolish(widget)
            widget.style().polish(widget)

    def _open_settings(self) -> None:
        dialog = CredentialDialog(self.config, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            values = dialog.values()
            if not all(values.get(key) for key in ("app_id", "api_key", "api_secret")):
                QMessageBox.warning(self, "缺少信息", "讯飞三项接口信息都必须填写。")
                return
            self.config.update(values)
            save_config(self.config)
            self._refresh_api_state()
            self.status_label.setText("接口设置已保存")

    def _sync_direction_note(self, direction: str) -> None:
        if direction == "auto":
            self.direction_note.setText("中英文混合识别 → 自动选择目标语言")
        elif direction == "zh_en":
            self.direction_note.setText("中文语音 → 讯飞英文译文")
        else:
            self.direction_note.setText("English speech → DeepSeek 中文译文")

    def _change_direction(self, direction: str) -> None:
        direction = direction if direction in {"auto", "zh_en", "en_zh"} else "zh_en"
        if direction in {"auto", "en_zh"} and not str(self.config.get("deepseek_api_key", "")).strip():
            QMessageBox.warning(self, "缺少 DeepSeek API Key", "请先在接口设置中填写 DeepSeek API Key。")
            self._apply_config_to_ui()
            return
        if direction == self.config.get("translation_direction"):
            return
        self.config["translation_direction"] = direction
        self._sync_direction_note(direction)
        self.auto_button.setChecked(direction == "auto")
        self.zh_en_button.setChecked(direction == "zh_en")
        self.en_zh_button.setChecked(direction == "en_zh")
        if self.overlay is not None:
            self.overlay.set_direction(direction)
        save_config(self.config)
        if self.client is not None:
            self.status_label.setText("正在重新连接麦克风…")
            old = self.client
            self.client = None
            self.client_generation += 1
            old.stop()
            QTimer.singleShot(1500, self._launch_client)

    def _update_materials_status(self) -> None:
        entries = self.config.get("glossary_entries", [])
        glossary_count = len(entries) if isinstance(entries, list) else 0
        reference_text = str(self.config.get("reference_text", "") or "")
        glossary_name = str(self.config.get("glossary_file_name", "") or "")
        reference_name = str(self.config.get("reference_file_name", "") or "")

        parts: list[str] = []
        if glossary_count:
            label = f"固定译法 {glossary_count} 条"
            if glossary_name:
                label += f"（{glossary_name}）"
            parts.append(label)
        if reference_text:
            label = f"参考稿 {len(reference_text)} 字"
            if reference_name:
                label += f"（{reference_name}）"
            parts.append(label)

        self.materials_status.setText(" · ".join(parts) if parts else "未导入固定翻译或参考稿件")

    def _sync_translation_materials_to_client(self) -> None:
        if self.client is not None:
            self.client.update_translation_materials(
                self.config.get("glossary_entries", []),
                str(self.config.get("reference_text", "") or ""),
            )

    def _import_glossary(self) -> None:
        selected, _ = QFileDialog.getOpenFileName(
            self,
            "上传固定翻译",
            str(Path.home() / "Documents"),
            "固定翻译 (*.txt *.csv *.tsv *.xlsx)",
        )
        if not selected:
            return
        try:
            entries = load_glossary_file(Path(selected))
        except Exception as exc:
            QMessageBox.critical(self, "导入失败", str(exc))
            return

        self.config["glossary_entries"] = entries
        self.config["glossary_file_name"] = Path(selected).name
        save_config(self.config)
        self._update_materials_status()
        self._sync_translation_materials_to_client()
        QMessageBox.information(
            self,
            "固定翻译已导入",
            f"已导入 {len(entries)} 条固定译法。\n\n"
            "示例：BIOEFFECT → 蓓欧菲\n"
            "英译中时 DeepSeek 会优先且固定使用右侧译法；"
            "中译英在启用 DeepSeek 资料翻译时也会使用对应译法。",
        )

    def _import_reference(self) -> None:
        selected, _ = QFileDialog.getOpenFileName(
            self,
            "导入参考稿件",
            str(Path.home() / "Documents"),
            "参考稿件 (*.txt *.docx *.csv *.tsv *.xlsx)",
        )
        if not selected:
            return
        try:
            reference_text = load_reference_file(Path(selected))
        except Exception as exc:
            QMessageBox.critical(self, "导入失败", str(exc))
            return

        self.config["reference_text"] = reference_text
        self.config["reference_file_name"] = Path(selected).name
        save_config(self.config)
        self._update_materials_status()
        self._sync_translation_materials_to_client()
        suffix = ""
        if len(reference_text) >= MAX_REFERENCE_CHARS:
            suffix = f"\n\n稿件较长，已使用前 {MAX_REFERENCE_CHARS} 个字符。"
        QMessageBox.information(
            self,
            "参考稿件已导入",
            f"已导入 {len(reference_text)} 个字符。翻译时会自动选取相关段落作为上下文。{suffix}",
        )
        self._start_reference_prewarm()

    def _start_reference_prewarm(self) -> None:
        """上传稿件后立即在后台预翻译，识别命中稿件时可直接取用预译文。"""
        api_key = str(self.config.get("deepseek_api_key", "")).strip()
        reference_text = str(self.config.get("reference_text", "") or "").strip()
        if not api_key or not reference_text:
            return
        from xfyun_client import DeepSeekTranslator

        translator = DeepSeekTranslator(
            api_key,
            str(self.config.get("deepseek_model", "deepseek-v4-flash")),
            glossary_entries=self.config.get("glossary_entries", []),
            reference_text=reference_text,
        )

        def emit(text: str) -> None:
            # 传译进行中不打扰主界面状态，静默完成。
            if self.client is None:
                self.events.put({"type": "status", "text": text, "state": "translating"})

        def on_progress(done: int, total: int) -> None:
            if done == total:
                emit(f"参考稿预翻译完成（{total} 段）")

        def run() -> None:
            emit("正在后台预翻译参考稿…")
            translator.prewarm_reference(on_progress=on_progress)

        threading.Thread(target=run, daemon=True).start()

    def _clear_translation_materials(self) -> None:
        has_materials = bool(
            self.config.get("glossary_entries") or self.config.get("reference_text")
        )
        if not has_materials:
            QMessageBox.information(self, "没有资料", "当前没有已导入的固定翻译或参考稿件。")
            return
        answer = QMessageBox.question(
            self,
            "清除翻译资料",
            "确定清除全部固定翻译和参考稿件吗？",
        )
        if answer != QMessageBox.StandardButton.Yes:
            return

        self.config["glossary_entries"] = []
        self.config["glossary_file_name"] = ""
        self.config["reference_text"] = ""
        self.config["reference_file_name"] = ""
        save_config(self.config)
        self._update_materials_status()
        self._sync_translation_materials_to_client()
        self.status_label.setText("翻译资料已清除")

    def _refresh_devices(self) -> None:
        self.status_label.setText("正在读取麦克风设备…")
        self.device_combo.setEnabled(False)

        def worker() -> None:
            try:
                self.events.put({"type": "devices", "devices": XfyunInterpreter.list_input_devices()})
            except Exception as exc:
                self.events.put({"type": "device_error", "text": str(exc)})

        threading.Thread(target=worker, daemon=True).start()

    def _selected_device(self) -> int | None:
        return self.device_combo.currentData()

    def _start(self) -> None:
        if not self._xfyun_ready():
            QMessageBox.warning(self, "接口未配置", "请先填写讯飞接口信息。")
            self._open_settings()
            return
        if self.config.get("translation_direction") in {"auto", "en_zh"} and not str(self.config.get("deepseek_api_key", "")).strip():
            QMessageBox.warning(self, "接口未配置", "自动识别和英译中需要 DeepSeek API Key。")
            self._open_settings()
            return

        self.session_segments = []
        self.session_chinese = ""
        self.session_english = ""
        self.session_started_at = datetime.now()
        self.session_finished_at = None
        self._ending_session = False
        self._service_retry_count = 0
        if self.result_window is not None:
            self.result_window.hide()

        overlay = self._ensure_overlay()
        overlay.apply_config(self.config)
        overlay.set_subtitles("", "")
        overlay.set_status("正在连接，请开始讲话…")
        overlay.show_for_translation()
        self.status_label.setText("正在启动传译…")
        self.hide()
        self._launch_client()

    def _launch_client(self) -> None:
        self.client_generation += 1
        session_id = self.client_generation
        try:
            self.client = XfyunInterpreter(
                app_id=str(self.config.get("app_id", "")),
                api_key=str(self.config.get("api_key", "")),
                api_secret=str(self.config.get("api_secret", "")),
                direction=str(self.config.get("translation_direction", "zh_en")),
                deepseek_api_key=str(self.config.get("deepseek_api_key", "")),
                deepseek_model=str(self.config.get("deepseek_model", "deepseek-v4-flash")),
                glossary_entries=self.config.get("glossary_entries", []),
                reference_text=str(self.config.get("reference_text", "") or ""),
                translation_interval_mode=str(
                    self.config.get("translation_interval_mode", "adaptive")
                ),
                play_tts=bool(self.config.get("play_tts", False)),
                input_device=self._selected_device(),
                on_event=self.events.put,
                session_id=session_id,
            )
            self.client.start()
            self.start_button.setEnabled(False)
            self.stop_button.setEnabled(True)
            self.wave.set_running(True)
        except Exception as exc:
            self.client = None
            append_log(f"启动失败：{exc}")
            if self.overlay:
                self.overlay.hide()
            self.show()
            self.raise_()
            self.activateWindow()
            QMessageBox.critical(self, "启动失败", str(exc))

    def _stop(self) -> None:
        if self._ending_session:
            return
        self._ending_session = True
        self.session_finished_at = datetime.now()
        self.start_button.setEnabled(False)
        self.stop_button.setEnabled(False)
        self.wave.set_running(False)
        if self.overlay is not None:
            self.config.update(self.overlay.current_settings())
            self.overlay.set_status("正在整理传译文档，请稍候…")
        save_config(self.config)
        self.status_label.setText("正在整理传译文档…")

        active_client = self.client
        generation = self.client_generation
        if active_client is None:
            if self.overlay is not None:
                self.overlay.hide()
            self._ending_session = False
            self.start_button.setEnabled(True)
            self._show_result_document()
            return

        def worker() -> None:
            try:
                result = active_client.finalize_session(timeout=20.0)
                self.events.put({"type": "session_finalized", "session_id": generation, **result})
            except Exception as exc:
                self.events.put({"type": "session_finalize_error", "session_id": generation, "text": str(exc)})

        self._finalize_thread = threading.Thread(target=worker, daemon=True)
        self._finalize_thread.start()

    def _tts_changed(self, enabled: bool) -> None:
        self.config["play_tts"] = bool(enabled)
        self.play_tts.blockSignals(True)
        self.play_tts.setChecked(bool(enabled))
        self.play_tts.blockSignals(False)
        if self.overlay is not None:
            self.overlay.set_play_tts(enabled)
        if self.client is not None:
            self.client.set_play_tts(enabled)
        self.save_timer.start(250)

    def _show_chinese_changed(self, enabled: bool) -> None:
        # 兼容旧配置：开启对应双语，关闭对应仅译文。
        self._set_overlay_mode("bilingual" if enabled else "translation")

    def _set_overlay_mode(self, mode: str) -> None:
        mode = mode if mode in {"source", "translation", "bilingual"} else "translation"
        self.config["overlay_display_mode"] = mode
        index = self.overlay_mode_combo.findData(mode)
        self.overlay_mode_combo.blockSignals(True)
        self.overlay_mode_combo.setCurrentIndex(max(0, index))
        self.overlay_mode_combo.blockSignals(False)
        if self.overlay is not None:
            self.overlay.apply_config(self.config)
        self.save_timer.start(250)

    def _overlay_mode_changed(self, _index: int) -> None:
        self._set_overlay_mode(str(self.overlay_mode_combo.currentData() or "translation"))

    def _translation_timing_changed(self, value) -> None:
        if isinstance(value, int):
            mode = str(self.timing_combo.currentData() or "adaptive")
        else:
            mode = str(value or "adaptive")
        if mode not in {"adaptive", "2", "3", "4", "5", "sentence"}:
            mode = "adaptive"
        self.config["translation_interval_mode"] = mode
        index = self.timing_combo.findData(mode)
        self.timing_combo.blockSignals(True)
        self.timing_combo.setCurrentIndex(max(0, index))
        self.timing_combo.blockSignals(False)
        if self.overlay is not None:
            self.overlay.set_timing_mode(mode)
        if self.client is not None:
            self.client.set_translation_interval_mode(mode)
        self.save_timer.start(250)

    def _overlay_settings_changed(self, values: dict) -> None:
        self.config.update(values)
        self.play_tts.blockSignals(True)
        self.play_tts.setChecked(bool(values.get("play_tts", False)))
        self.play_tts.blockSignals(False)
        self._set_overlay_mode(str(values.get("overlay_display_mode", "translation")))
        self._translation_timing_changed(
            str(values.get("translation_interval_mode", "adaptive"))
        )
        self.save_timer.start(350)

    def _poll_events(self) -> None:
        while True:
            try:
                event = self.events.get_nowait()
            except queue.Empty:
                return
            session_id = event.get("session_id")
            if session_id is not None and int(session_id) != self.client_generation:
                continue
            event_type = event.get("type")
            if event_type == "devices":
                self.device_combo.clear()
                self.device_combo.addItem("系统默认麦克风（推荐）", None)
                for index, name in event.get("devices", []):
                    self.device_combo.addItem(f"{index}: {name}", index)
                self.device_combo.setEnabled(True)
                self.status_label.setText("麦克风设备已刷新")
            elif event_type == "device_error":
                self.device_combo.setEnabled(True)
                self.status_label.setText("设备读取失败，可使用系统默认麦克风")
            elif event_type == "status":
                text = str(event.get("text", ""))
                state = str(event.get("state", ""))
                if state == "connected":
                    self._service_retry_count = 0
                self.status_label.setText(text)
                if self.overlay is not None:
                    self.overlay.set_status(text)
            elif event_type == "subtitles":
                chinese = str(event.get("chinese", ""))
                english = str(event.get("english", ""))
                segments = event.get("segments", [])
                self.session_chinese = chinese
                self.session_english = english
                if isinstance(segments, list):
                    self.session_segments = [
                        dict(segment) for segment in segments if isinstance(segment, dict)
                    ]
                if self.overlay is not None:
                    self.overlay.set_subtitles(
                        chinese,
                        english,
                        str(event.get("source_language", "")) or None,
                        str(event.get("source_transcript", "")),
                        str(event.get("translation_transcript", "")),
                    )
                if self.result_window is not None and self.result_window.isVisible():
                    self.result_window.set_document(
                        self.session_segments,
                        self.session_chinese,
                        self.session_english,
                        self.session_started_at,
                        self.session_finished_at or datetime.now(),
                    )
            elif event_type == "finalize_progress":
                completed = int(event.get("completed", 0) or 0)
                total = int(event.get("total", 0) or 0)
                text = str(event.get("text", f"正在整理传译文档：{completed}/{total}"))
                self.status_label.setText(text)
                if self.overlay is not None:
                    self.overlay.set_status(text)
            elif event_type == "session_finalized":
                segments = event.get("segments", [])
                if isinstance(segments, list):
                    self.session_segments = [dict(s) for s in segments if isinstance(s, dict)]
                self.session_chinese = str(event.get("chinese", ""))
                self.session_english = str(event.get("english", ""))
                self.client = None
                self._ending_session = False
                self.start_button.setEnabled(True)
                self.stop_button.setEnabled(False)
                if self.overlay is not None:
                    self.overlay.hide()
                self.status_label.setText("传译文档整理完成")
                self._show_result_document()
            elif event_type == "session_finalize_error":
                text = str(event.get("text", "文档整理失败"))
                append_log(f"文档整理失败：{text}")
                self.client = None
                self._ending_session = False
                self.start_button.setEnabled(True)
                if self.overlay is not None:
                    self.overlay.hide()
                self._show_result_document()
                QMessageBox.warning(self.result_window or self, "文档整理未完全完成", text)
            elif event_type == "protocol_warning":
                text = str(event.get("text", "收到异常返回，已继续监听"))
                detail = str(event.get("detail", ""))
                append_log(
                    f"{text}" + (f" | {detail}" if detail else "")
                )
                self.status_label.setText(text)
                if self.overlay is not None:
                    self.overlay.set_status(text)
                # 不弹出对话框、不停止麦克风、不返回主界面。
                continue
            elif event_type == "deepseek_error":
                text = str(event.get("text", "DeepSeek 翻译失败"))
                self.status_label.setText("英译中失败")
                if self.overlay is not None:
                    self.overlay.set_status("English 已识别 · DeepSeek 翻译失败")
                append_log(text)
            elif event_type == "error":
                text = str(event.get("text", "未知错误"))
                code = int(event.get("code", 0) or 0)
                append_log(text)
                if code == 10008 and not self._ending_session and self._service_retry_count < 3:
                    self._service_retry_count += 1
                    attempt = self._service_retry_count
                    wait_ms = 1500 * attempt
                    self.client = None
                    # 立即提高 generation，使旧连接后续的 closed 事件失效。
                    self.client_generation += 1
                    message = f"讯飞服务实例暂时不可用，正在自动重连（{attempt}/3）…"
                    self.status_label.setText(message)
                    if self.overlay is not None:
                        self.overlay.set_status(message)
                    QTimer.singleShot(wait_ms, self._launch_client)
                    continue
                if code == 10008:
                    friendly = (
                        "讯飞服务实例暂时不可用，软件已经自动重试 3 次。\n\n"
                        "这通常属于讯飞服务容量或服务实例状态问题，并非麦克风故障。"
                        "请稍后重新开始；若持续出现，请在讯飞控制台确认同声传译服务状态和授权。"
                    )
                else:
                    friendly = text
                self._abort_session_to_settings("传译错误", friendly)
            elif event_type == "closed":
                self.wave.set_running(False)
                if self._ending_session:
                    self.status_label.setText("连接已关闭，正在完成剩余译文…")
                    if self.overlay is not None:
                        self.overlay.set_status("连接已关闭，正在完成剩余译文…")
                    continue
                self.client = None
                self.start_button.setEnabled(True)
                self.stop_button.setEnabled(False)
                self.status_label.setText(str(event.get("text", "连接已关闭")))
                if self.overlay is not None and self.overlay.isVisible():
                    self.overlay.set_status("连接已关闭，请点击停止返回")


    def closeEvent(self, event: QCloseEvent) -> None:  # type: ignore[override]
        if self.client is not None:
            self.client.stop()
        if self.overlay is not None:
            self.config.update(self.overlay.current_settings())
            self.overlay.close()
        if self.result_window is not None:
            self.result_window.hide()
        save_config(self.config)
        event.accept()


def stylesheet() -> str:
    chevron_url = QUrl.fromLocalFile(str(resource_path("assets/combo_chevron.png"))).toString()
    return f"""
    * {{ font-family: -apple-system, BlinkMacSystemFont, 'SF Pro Display', 'Microsoft YaHei', 'Segoe UI', sans-serif; }}
    QMainWindow, QWidget {{ background: {TOKENS['bg']}; color: {TOKENS['text']}; }}
    QMainWindow#subtitleOverlay, QWidget#overlayRoot, QWidget#subtitleText, QSplitter#subtitleSplitter {{ background: transparent; }}
    QLabel#versionLabel {{ color: {TOKENS['muted']}; font-size: 12px; font-weight: 600; }}
    QLabel#heroStatus {{ color: {TOKENS['text']}; font-size: 23px; font-weight: 700; }}
    QLabel#settingsTitle {{ color: {TOKENS['muted']}; font-size: 12px; font-weight: 700; }}
    QLabel#settingsLabel {{ color: {TOKENS['muted']}; font-size: 13px; min-width: 56px; }}
    QFrame#hairline {{ background: {TOKENS['border']}; border: none; max-height: 1px; }}
    QLabel#dialogTitle {{ color: {TOKENS['deep_blue']}; font-size: 24px; font-weight: 700; }}
    QLabel#mutedLabel, QLabel#flowNote, QLabel#statusLabel {{ color: {TOKENS['muted']}; }}
    QLabel#statusDot[state='ok'] {{ color: {TOKENS['green']}; font-size: 16px; }}
    QLabel#statusDot[state='partial'] {{ color: #D9A928; font-size: 16px; }}
    QLabel#statusDot[state='error'] {{ color: {TOKENS['danger']}; font-size: 16px; }}
    QLabel#apiState[state='ok'] {{ color: {TOKENS['green']}; font-weight: 650; }}
    QLabel#apiState[state='partial'] {{ color: #997415; font-weight: 650; }}
    QLabel#apiState[state='error'] {{ color: {TOKENS['danger']}; font-weight: 650; }}
    QMainWindow#documentWindow, QWidget#documentBody {{ background: #F4F7FB; }}
    QLabel#documentTitle {{ color: {TOKENS['deep_blue']}; font-size: 32px; font-weight: 750; }}
    QLabel#documentMeta {{ color: {TOKENS['muted']}; font-size: 13px; }}
    QFrame#documentColumnHeader {{ background: #E8F0FA; border: 1px solid #CBDCEE; border-radius: 10px; }}
    QLabel#documentColumnTitle {{ color: {TOKENS['deep_blue']}; font-size: 15px; font-weight: 700; }}
    QScrollArea#documentScroll {{ background: transparent; border: none; }}
    QFrame#documentBlock {{ background: white; border: 1px solid {TOKENS['border']}; border-radius: 12px; }}
    QLabel#documentNumber {{ color: {TOKENS['blue']}; font-size: 13px; font-weight: 750; padding-top: 9px; }}
    QFrame#documentColumn {{ background: #FBFCFE; border: 1px solid #E3EAF3; border-radius: 9px; }}
    QLabel#documentLanguageLabel {{ color: {TOKENS['muted']}; font-size: 11px; font-weight: 700; }}
    QLabel#documentParagraph {{ color: {TOKENS['text']}; font-size: 17px; line-height: 1.65; }}
    QLabel#documentMissing {{ color: {TOKENS['danger']}; font-size: 16px; font-weight: 650; }}
    QLabel#documentEmpty {{ color: {TOKENS['muted']}; font-size: 16px; padding: 80px; }}
    QPushButton {{ min-height: 38px; padding: 0 16px; border-radius: 9px; border: 1px solid {TOKENS['border']}; background: {TOKENS['white']}; color: {TOKENS['deep_blue']}; font-weight: 600; }}
    QPushButton:hover {{ background: {TOKENS['light_blue']}; border-color: #AEC5E5; }}
    QPushButton:pressed {{ background: #DDEAF9; }}
    QPushButton:disabled {{ background: {TOKENS['disabled']}; color: #A4ADBA; border-color: #E0E5EC; }}
    QPushButton#startButton {{ min-height: 58px; max-height: 58px; background: {TOKENS['blue']}; color: white; border: none; border-radius: 14px; font-size: 18px; font-weight: 750; }}
    QPushButton#startButton:hover {{ background: #0F5EC5; }}
    QPushButton#startButton:pressed {{ background: {TOKENS['deep_blue']}; }}
    QPushButton#startButton:disabled {{ background: {TOKENS['disabled']}; color: #A4ADBA; }}
    QPushButton#stopButton {{ min-height: 56px; max-height: 56px; background: white; color: {TOKENS['danger']}; border: 1px solid #E7B6BB; border-radius: 14px; font-size: 15px; font-weight: 650; min-width: 108px; }}
    QPushButton#stopButton:hover {{ background: #FDF0F1; border-color: {TOKENS['danger']}; }}
    QPushButton#stopButton:pressed {{ background: #F7D9DC; }}
    QPushButton#stopButton:disabled {{ background: {TOKENS['disabled']}; color: #A4ADBA; border-color: #E0E5EC; }}
    QPushButton#directionButton {{ min-width: 108px; min-height: 38px; max-height: 38px; border-radius: 11px; background: white; color: {TOKENS['text']}; border: 1px solid {TOKENS['border']}; font-weight: 600; }}
    QPushButton#directionButton:hover {{ background: {TOKENS['light_blue']}; }}
    QPushButton#directionButton:pressed {{ background: #DDEAF9; }}
    QPushButton#directionButton:checked {{ background: {TOKENS['blue']}; color: white; border-color: {TOKENS['blue']}; font-weight: 700; }}
    QPushButton#primaryButton {{ background: {TOKENS['blue']}; color: white; border: none; padding: 0 24px; }}
    QPushButton#primaryButton:hover {{ background: #0F5EC5; }}
    QPushButton#primaryButton:pressed {{ background: {TOKENS['deep_blue']}; }}
    QPushButton#outlineButton {{ background: white; color: {TOKENS['deep_blue']}; border: 1px solid #9DB6D8; }}
    QPushButton#ghostButton {{ background: transparent; border: none; color: {TOKENS['muted']}; min-height: 34px; padding: 0 10px; }}
    QPushButton#ghostButton:hover {{ background: transparent; color: {TOKENS['danger']}; }}
    QLineEdit, QSpinBox, QTextEdit {{ background: white; color: {TOKENS['text']}; border: 1px solid {TOKENS['border']}; border-radius: 8px; padding: 8px 12px; selection-background-color: {TOKENS['blue']}; }}
    QLineEdit:focus, QSpinBox:focus, QTextEdit:focus {{ border: 2px solid #90B8EA; }}
    QLineEdit, QSpinBox {{ min-height: 38px; }}
    QTextEdit {{ line-height: 1.55; }}
    QComboBox {{ background: white; color: {TOKENS['text']}; border: 1px solid {TOKENS['border']}; border-radius: 8px; padding: 3px 10px; padding-right: 26px; min-height: 34px; selection-background-color: {TOKENS['blue']}; }}
    QComboBox:hover {{ border-color: #AEC5E5; }}
    QComboBox:focus {{ border: 1px solid #90B8EA; }}
    QComboBox::drop-down {{ subcontrol-origin: padding; subcontrol-position: center right; width: 24px; border: none; background: transparent; }}
    QComboBox::down-arrow {{ image: url({chevron_url}); width: 12px; height: 8px; }}
    QComboBox QAbstractItemView {{ background: white; border: 1px solid {TOKENS['border']}; border-radius: 10px; padding: 5px; outline: 0; selection-background-color: transparent; }}
    QComboBox QAbstractItemView::item {{ min-height: 26px; padding: 0 10px; border-radius: 6px; color: {TOKENS['text']}; }}
    QComboBox QAbstractItemView::item:hover, QComboBox QAbstractItemView::item:selected {{ background: {TOKENS['light_blue']}; color: {TOKENS['deep_blue']}; }}
    QCheckBox {{ spacing: 8px; color: {TOKENS['text']}; }}
    QCheckBox::indicator {{ width: 18px; height: 18px; border: 1px solid #AFC0D5; border-radius: 4px; background: white; }}
    QCheckBox::indicator:checked {{ background: {TOKENS['blue']}; border-color: {TOKENS['blue']}; }}
    QSplitter#subtitleSplitter::handle:vertical {{ height: 10px; background: rgba(23,104,213,155); margin-left: 18px; margin-right: 18px; border-radius: 5px; }}
    QSplitter#subtitleSplitter::handle:vertical:hover {{ background: rgba(23,104,213,225); }}
    QFrame#overlaySideToolbar QPushButton {{ min-height: 32px; max-height: 34px; min-width: 48px; padding: 0 4px; border-radius: 8px; }}
    QPushButton#overlayModeButton:checked {{ background: {TOKENS['blue']}; color: white; border-color: {TOKENS['blue']}; }}
    QPushButton#overlayToolButton {{ background: rgba(255,255,255,205); }}
    QPushButton#overlayStop {{ background: {TOKENS['deep_blue']}; color: white; border: none; }}
    QLabel#overlayStatus {{ color: {TOKENS['muted']}; font-size: 12px; }}
    QSizeGrip#overlayResizeGrip {{ background: transparent; }}
    """


def install_exception_hook() -> None:
    original = sys.excepthook

    def handler(exc_type, exc_value, exc_traceback) -> None:
        details = "".join(traceback.format_exception(exc_type, exc_value, exc_traceback))
        append_log("未捕获异常：\n" + details)
        try:
            QMessageBox.critical(None, "程序错误", f"程序遇到错误：\n{exc_value}\n\n日志：{LOG_FILE}")
        except Exception:
            pass
        original(exc_type, exc_value, exc_traceback)

    sys.excepthook = handler


def main() -> int:
    app = QApplication(sys.argv)
    app.setApplicationName(APP_NAME)
    app.setApplicationVersion(APP_VERSION)
    app.setOrganizationName("ismolar")
    app.setStyle("Fusion")
    app.setStyleSheet(stylesheet())
    try:
        # 翻译记忆磁盘持久化：内存热缓存满时落盘，重启后记忆保留。
        set_manuscript_cache_db(CONFIG_DIR / "manuscript_memory.db")
    except Exception:
        pass
    icon_path = resource_path("assets/app_icon.png")
    if icon_path.exists():
        app.setWindowIcon(QIcon(str(icon_path)))
    install_exception_hook()
    window = MainWindow()
    window.show()
    return app.exec()


if __name__ == "__main__":
    raise SystemExit(main())
